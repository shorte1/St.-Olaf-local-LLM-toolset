import pdfplumber
from docx import Document
import os


def extract_text(filepath):
    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.txt':
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()

    elif ext == '.pdf':
        text = ''
        page_count = 0
        with pdfplumber.open(filepath) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + '\n\n'
        text = text.strip()

        # Fewer than 50 chars per page almost certainly means a scanned PDF
        if page_count > 0 and len(text) < page_count * 50:
            text = _ocr_pdf(filepath)

        return text

    elif ext == '.docx':
        doc = Document(filepath)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        return '\n\n'.join(paragraphs)

    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _ocr_pdf(filepath):
    try:
        import pypdfium2 as pdfium
        import pytesseract
    except ImportError as e:
        raise ValueError(
            "This PDF appears to be a scan with no embedded text. "
            "OCR support requires pytesseract and tesseract.\n"
            "Install with: conda install -c conda-forge tesseract pytesseract"
        ) from e

    try:
        pdf = pdfium.PdfDocument(filepath)
    except Exception as e:
        raise ValueError(f"Could not open PDF for OCR: {e}") from e

    text = ''
    for page in pdf:
        # scale=2 gives 144 DPI which is enough for clean OCR
        bitmap = page.render(scale=2)
        pil_image = bitmap.to_pil()
        try:
            page_text = pytesseract.image_to_string(pil_image)
        except pytesseract.TesseractNotFoundError:
            raise ValueError(
                "Tesseract binary not found. "
                "Install with: conda install -c conda-forge tesseract"
            )
        if page_text:
            text += page_text + '\n\n'

    return text.strip()
